from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "STRH_HRIS_Initial_Testing_User_Manual.docx"
LOGO = ROOT / "STRH-logo.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.line_spacing = 1.15


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=INK)
        add_run(p, text[len(bold_prefix) :])
    else:
        add_run(p, text)
    return p


def add_bullet(doc, text):
    return add_para(doc, text, style="List Bullet")


def add_number(doc, text):
    return add_para(doc, text, style="List Number")


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_margins(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    add_run(p, f"{title}: ", bold=True, color=DARK_BLUE)
    add_run(p, body)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_shading(hdr[idx], header_fill)
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, h, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            add_run(p, str(value))
    doc.add_paragraph()
    return table


def add_checklist_table(doc, rows):
    add_table(
        doc,
        ["Done", "Test Activity", "Expected Result / What To Record"],
        rows,
        [900, 3600, 4860],
    )


def setup_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "STRH HRIS Initial Testing User Manual", color=MUTED, size=9)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "For STRH HR initial testing only | Prepared June 2026", color=MUTED, size=9)


def build_manual():
    doc = Document()
    style_doc(doc)
    setup_header_footer(doc)

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.0))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(kicker, "Initial Testing Guide", bold=True, color=BLUE, size=11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "STRH HRIS User Manual", bold=True, color=INK, size=26)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        subtitle,
        "A practical guide for HR personnel testing the Human Resource Information System",
        color=MUTED,
        size=13,
    )

    meta = [
        ("Audience", "STRH HR testers, HR supervisors, and system coordinators"),
        ("Testing Stage", "Initial user testing / UAT familiarization"),
        ("System URL", "To be provided after port forwarding or local network deployment"),
        ("Recommended Browser", "Current Google Chrome or Microsoft Edge"),
    ]
    add_table(doc, ["Field", "Details"], meta, [2700, 6660])

    add_note(
        doc,
        "Purpose",
        "This manual helps HR test the current system behavior and record issues clearly. It is not yet the final production operations manual.",
    )

    doc.add_page_break()

    add_para(doc, "Contents", "Heading 1")
    for item in [
        "1. Testing Preparation",
        "2. Sign In and Navigation",
        "3. User Roles",
        "4. Module Quick Reference",
        "5. HR Testing Scripts",
        "6. Known Initial Testing Notes",
        "7. Issue Reporting Template",
    ]:
        add_bullet(doc, item)

    add_para(doc, "1. Testing Preparation", "Heading 1")
    add_para(
        doc,
        "Before testing, confirm that the system is running and that HR has been given test accounts. For local development, the project launcher starts the application at http://localhost:47100. For STRH testing through port forwarding, use the forwarded URL provided by the system administrator.",
    )
    add_checklist_table(
        doc,
        [
            ("[ ]", "Confirm the test URL opens.", "Login page appears with STRH/agency branding."),
            ("[ ]", "Confirm HR test account credentials.", "Tester can sign in and see HR modules."),
            ("[ ]", "Prepare sample employee records.", "Use real-like but approved test data only."),
            ("[ ]", "Prepare feedback notes.", "Record exact module, action, expected result, actual result, and screenshot if possible."),
        ],
    )
    add_note(
        doc,
        "Data Privacy",
        "Use only authorized test data. Do not upload confidential records unless STRH has approved the testing environment for that data.",
    )

    add_para(doc, "2. Sign In and Navigation", "Heading 1")
    add_number(doc, "Open the system URL in a supported browser.")
    add_number(doc, "Enter the assigned username and password.")
    add_number(doc, "Click Sign In.")
    add_number(doc, "If the account is using a temporary password, change the password when prompted.")
    add_number(doc, "Use the left sidebar on desktop or bottom navigation on mobile to move between modules.")
    add_para(
        doc,
        "The dashboard shows workforce totals, employment type mix, age profile, gender distribution, employees by division, and quick links for common HR actions.",
    )

    add_para(doc, "3. User Roles", "Heading 1")
    add_table(
        doc,
        ["Role", "Typical Access During Testing"],
        [
            ("Admin", "System administration, users, backups, audit/error logs, plus HR modules."),
            ("HR", "Dashboard, employee records, attendance, plantilla, movements, service records, leave, reports, and settings."),
            ("Viewer", "Read-oriented access to dashboard, employees, plantilla, movements, service records, and reports."),
            ("Employee", "Self-service dashboard, profile, attendance, leave filing, requests, and PDS generation."),
        ],
        [1800, 7560],
    )
    add_note(
        doc,
        "Account Locking",
        "Three failed login attempts lock the account. An Admin can unlock the account or reset the password.",
    )

    add_para(doc, "4. Module Quick Reference", "Heading 1")
    add_table(
        doc,
        ["Module", "Use It To Test"],
        [
            ("Dashboard", "Verify summary counts, charts, and quick links."),
            ("Employee Management", "Search, filter, add employee, open 201 file, and remove a record from the active list."),
            ("201 File", "Update personal/employment details and maintain family, education, eligibility, work, training, service, and IPCR records."),
            ("Attendance DTR", "Import biometric/file punches, refresh DTR, add/edit/delete DTR rows, handle correction requests, and generate DTR Excel/PDF."),
            ("Plantilla & PSIPOP", "Create plantilla items, assign/vacate employees, view occupancy, vacancy, and movement history."),
            ("Employee Movements", "Prepare, submit, review, approve, post, reverse, and inspect action history for personnel movements."),
            ("Service Records", "Review movement-derived service history, add legacy periods, and export HRIS-generated Excel/PDF service records."),
            ("Leave Management", "File leave, approve/disapprove/cancel, generate CS Form No. 6 Excel/PDF, and manage leave types."),
            ("Self-Service / My Requests", "Validate employee-side leave filing, DTR correction requests, request history, and PDS generation."),
            ("Reports & Analytics", "Currently a report catalog; broad report generation is still pending."),
            ("System Administration", "Manage users, reset passwords, unlock accounts, create backups, and review audit/error logs."),
            ("Settings / References", "Update agency branding and maintain reference libraries such as departments, positions, salary grades, and HR coded values."),
        ],
        [2100, 7260],
    )

    add_para(doc, "5. HR Testing Scripts", "Heading 1")
    add_para(doc, "Employee Management and 201 File", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "Search by name, ID, or email.", "Correct employees appear; no unrelated record dominates the result."),
            ("[ ]", "Filter by department, employment type, status, and gender.", "Counts and list update correctly."),
            ("[ ]", "Add a test employee.", "New employee appears in the list and can be opened."),
            ("[ ]", "Open the 201 File.", "Personal, employment, contact, ID, and record tabs load."),
            ("[ ]", "Add one row in education/training/work/IPCR.", "Row saves, appears in the table, and can be edited/deleted."),
            ("[ ]", "Generate or preview PDS where available.", "File downloads or preview opens without layout error."),
        ],
    )

    add_para(doc, "Attendance and DTR", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "Filter by employee and date range.", "DTR rows match selected employee and period."),
            ("[ ]", "Import DTR from biometric device or uploaded file.", "Punches import and DTR rows refresh."),
            ("[ ]", "Add or edit a manual DTR row.", "Manual row saves and audit/correction behavior is preserved."),
            ("[ ]", "Submit a DTR correction request.", "Request appears as Pending with original and requested values."),
            ("[ ]", "Approve and apply correction.", "Requested values are applied; audit shows reviewer and remarks."),
            ("[ ]", "Generate DTR Excel/PDF.", "Export contains the selected employee, period, noter, and rows."),
        ],
    )

    add_para(doc, "Leave Management", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "File leave for an employee.", "Application is recorded with dates, days, type, reason, and Form 6 details."),
            ("[ ]", "Approve leave.", "Status changes to Approved and applicable leave credits are deducted."),
            ("[ ]", "Disapprove or cancel leave.", "Decision remarks are stored and status changes correctly."),
            ("[ ]", "Generate CS Form No. 6 Excel.", "Official workbook-based export downloads."),
            ("[ ]", "Generate CS Form No. 6 PDF preview.", "PDF opens or downloads; print layout should be reviewed by HR."),
            ("[ ]", "Review leave credit ledger.", "Manual adjustments, approvals, and reversals are traceable."),
        ],
    )

    add_para(doc, "Plantilla, Movements, and Service Records", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "Create or edit a plantilla item.", "Item number, position, salary/funding, and status save correctly."),
            ("[ ]", "Assign and vacate an item.", "Occupancy and vacancy status update without duplicate active occupancy."),
            ("[ ]", "Prepare a personnel movement.", "Draft is created with employee, action type, effective date, authority, and target details."),
            ("[ ]", "Submit, review, approve, and post a movement.", "Posting updates employee profile and plantilla occupancy."),
            ("[ ]", "Reverse a posted movement with reason.", "Prior employee/plantilla state is restored when allowed."),
            ("[ ]", "Open Service Records for the employee.", "Posted movement periods appear; legacy period can be added if needed."),
            ("[ ]", "Export service record Excel/PDF.", "Generated export is readable and reflects the encoded service record data."),
        ],
    )

    add_para(doc, "Administration, Settings, and References", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "Create a user and link to employee when needed.", "Temporary password is shown; account can log in."),
            ("[ ]", "Reset password and unlock user.", "Temporary password works and failed-login counter resets."),
            ("[ ]", "Create a manual backup.", "Backup file appears and can be downloaded by Admin."),
            ("[ ]", "Update agency profile, logo, icon, or cover photo.", "Login and app branding reflect saved settings."),
            ("[ ]", "Add/update/deactivate a reference value.", "Validation prevents invalid hierarchy, duplicate, or inactive-parent choices."),
        ],
    )

    add_para(doc, "Employee Self-Service Checks", "Heading 2")
    add_checklist_table(
        doc,
        [
            ("[ ]", "Log in as an Employee account.", "Employee sees My Dashboard, My Profile, Attendance, Requests, and Self-Service."),
            ("[ ]", "Apply for leave.", "Request appears in HR Leave Management and My Requests."),
            ("[ ]", "Request DTR correction.", "HR can review the request from Attendance."),
            ("[ ]", "Generate PDS.", "PDS file is generated for the linked employee."),
            ("[ ]", "Try unrelated employee access.", "Employee must not access another employee's confidential record."),
        ],
    )

    add_para(doc, "6. Known Initial Testing Notes", "Heading 1")
    for note in [
        "Reports & Analytics is currently a catalog; broad report generation/export is still pending.",
        "Service Record exports are generated directly from encoded HRIS service history and can be aligned with STRH's preferred final print format.",
        "Some self-service actions such as certificate request, overtime, schedule change, and similar HR requests are placeholders or future workflow items.",
        "Leave accrual automation, low-balance warnings, configurable multi-level approvals, automated backups, restore workflow, and archival/retention tools require later validation or implementation.",
        "Biometric testing must be done with the actual STRH devices, network, time settings, and employee device IDs.",
        "Port forwarding and production URL setup will be handled later; this manual only prepares HR for application testing.",
    ]:
        add_bullet(doc, note)

    add_para(doc, "7. Issue Reporting Template", "Heading 1")
    add_table(
        doc,
        ["Field", "What HR Should Write"],
        [
            ("Tester", "Name of person testing."),
            ("Date / Time", "When the issue happened."),
            ("Module", "Example: Leave Management, Attendance DTR, Employee 201 File."),
            ("Test Account / Role", "Example: HR, Admin, Employee."),
            ("Steps Taken", "Numbered actions performed before the issue occurred."),
            ("Expected Result", "What HR expected the system to do."),
            ("Actual Result", "What the system did instead."),
            ("Screenshot / File", "Attach screenshot or generated export when useful."),
            ("Severity", "Low, Medium, High, or Blocking."),
        ],
        [2400, 6960],
    )

    add_para(doc, "Suggested Testing Sign-Off", "Heading 2")
    add_table(
        doc,
        ["Review Area", "Reviewer", "Decision / Comments", "Date"],
        [
            ("Employee Records / 201 File", "", "", ""),
            ("Attendance and DTR", "", "", ""),
            ("Leave and Form 6", "", "", ""),
            ("Plantilla / Movements / Service Records", "", "", ""),
            ("Self-Service", "", "", ""),
            ("Administration / Security / Backup", "", "", ""),
        ],
        [2300, 1700, 3960, 1400],
    )

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_manual())
