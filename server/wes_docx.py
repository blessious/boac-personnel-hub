import json
import re
import sys
from copy import deepcopy
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ModuleNotFoundError:
    print(
        "python-docx is required for Work Experience Sheet export. Run setup.bat to install Python dependencies.",
        file=sys.stderr,
    )
    raise


FONT_NAME = "Arial"
FONT_SIZE = Pt(12)


def text(value):
    if value is None:
        return ""
    return str(value).strip()


def first(*values):
    for value in values:
        normalized = text(value)
        if normalized:
            return normalized
    return ""


def full_name(employee):
    parts = [
        employee.get("firstname"),
        employee.get("middlename"),
        employee.get("lastname"),
        employee.get("nameExt"),
    ]
    return " ".join(text(part) for part in parts if text(part))


def parse_date(value):
    value = text(value)
    if not value:
        return None
    if value.lower() in {"present", "current"}:
        return None
    normalized = value.replace("Z", "+00:00")
    for candidate in (normalized, value):
        try:
            return datetime.fromisoformat(candidate).date()
        except ValueError:
            pass
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(value, fmt).date()
        except ValueError:
            pass
    return None


def date_text(value, present=False):
    raw = text(value)
    if present and not raw:
        return "Present"
    if raw.lower() in {"present", "current"}:
        return "Present"
    parsed = parse_date(raw)
    if parsed:
        return f"{parsed.day} {parsed.strftime('%B %Y')}"
    return raw


def duration(item):
    date_from = date_text(first(item.get("dateFrom"), item.get("from")))
    date_to = date_text(first(item.get("dateTo"), item.get("to")), present=True)
    if date_from and date_to:
        return f"{date_from} – {date_to}"
    return date_from or date_to or "N/A"


def sort_current_first(rows):
    def key(item):
        date_to = text(item.get("dateTo") or item.get("to")).lower()
        if date_to in {"", "present", "current"}:
            return "9999-12-31"
        parsed = parse_date(date_to)
        if parsed:
            return parsed.isoformat()
        return date_to

    return sorted(rows, key=key, reverse=True)


def section_payloads(sections, key):
    return [row.get("payload", {}) for row in sections.get(key, []) if isinstance(row, dict)]


def clear_cell(cell):
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is not tc_pr:
            tc.remove(child)

    # Keep the required cell paragraph, but make it plain so old list
    # formatting from the sample template cannot create blank bullets.
    tc.append(OxmlElement("w:p"))


def replace_cell_paragraphs(cell, paragraph_elements):
    tc = cell._tc
    tc_pr = tc.tcPr
    for child in list(tc):
        if child is not tc_pr:
            tc.remove(child)
    for paragraph_element in paragraph_elements:
        tc.append(deepcopy(paragraph_element))


def apply_run_font(run):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT_NAME)


def set_paragraph_text(paragraph, value):
    """Replace text while retaining the template paragraph and run formatting."""
    source_run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            source_run_properties = deepcopy(run._r.rPr)
            break

    paragraph_properties = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not paragraph_properties:
            paragraph._p.remove(child)

    if value:
        run_element = OxmlElement("w:r")
        if source_run_properties is not None:
            run_element.append(source_run_properties)
        text_element = OxmlElement("w:t")
        if value[:1].isspace() or value[-1:].isspace() or "  " in value:
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = value
        run_element.append(text_element)
        paragraph._p.append(run_element)


def add_line(cell, label, value="", bold_label=True):
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    if label:
        run = paragraph.add_run(label)
        run.bold = bold_label
        apply_run_font(run)
    if value:
        apply_run_font(paragraph.add_run(value))
    return paragraph


def content_lines(value):
    values = []
    for raw_line in text(value).splitlines():
        line = raw_line.strip()
        line = re.sub(r"^(?:[•●◦▪\-–—]|o(?=\s)|\d+[.)])\s*", "", line, flags=re.IGNORECASE)
        if line:
            values.append(line)
    return values or ["N/A"]


def fill_work_row(cell, item, source_cell):
    """Build a record exclusively from paragraph archetypes in the WES template."""
    source = source_cell.paragraphs
    if len(source) < 16:
        raise RuntimeError("Work Experience Sheet template record layout is incomplete")

    paragraph_elements = [
        source[0]._p,
        source[1]._p,
        source[2]._p,
        source[3]._p,
        source[4]._p,
        source[5]._p,
        source[6]._p,
        source[7]._p,
    ]
    paragraph_elements.extend(source[9]._p for _ in content_lines(item.get("accomplishments")))
    paragraph_elements.extend([source[11]._p, source[12]._p, source[13]._p])
    paragraph_elements.extend(source[14]._p for _ in content_lines(item.get("actualDuties")))
    paragraph_elements.append(source[15]._p)
    replace_cell_paragraphs(cell, paragraph_elements)

    paragraphs = cell.paragraphs
    set_paragraph_text(paragraphs[1], f"Duration:  {duration(item)}")
    set_paragraph_text(
        paragraphs[2], f"Position:  {first(item.get('position'), item.get('designation')) or 'N/A'}"
    )
    set_paragraph_text(paragraphs[3], f"Name of Office/Unit: {text(item.get('officeUnit')) or 'N/A'}")
    set_paragraph_text(
        paragraphs[4], f"Immediate Supervisor: {text(item.get('immediateSupervisor')) or 'N/A'}"
    )
    set_paragraph_text(
        paragraphs[5],
        " Name of Agency/Organization and Location: "
        + (first(item.get("agencyOrganizationLocation"), item.get("company")) or "N/A"),
    )

    accomplishments = content_lines(item.get("accomplishments"))
    first_accomplishment = 8
    for offset, value in enumerate(accomplishments):
        set_paragraph_text(paragraphs[first_accomplishment + offset], value)

    summary_heading = first_accomplishment + len(accomplishments) + 1
    duties = content_lines(item.get("actualDuties"))
    first_duty = summary_heading + 2
    for offset, value in enumerate(duties):
        set_paragraph_text(paragraphs[first_duty + offset], value)


def clone_row(table, source_row_element):
    row_element = deepcopy(source_row_element)
    table._tbl.append(row_element)
    return table.rows[-1]


def fill_docx(payload, output_path, template_path):
    document = Document(template_path)
    employee = payload.get("employee", {})
    sections = payload.get("sections", {})
    rows = sort_current_first(section_payloads(sections, "work"))

    if not document.tables:
        raise RuntimeError("Work Experience Sheet template does not contain a table")

    table = document.tables[0]
    template_row = table.rows[2] if len(table.rows) > 2 else table.rows[-1]
    template_row_element = deepcopy(template_row._tr)
    template_cell = template_row.cells[0]
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)

    if not rows:
        row = clone_row(table, template_row_element)
        clear_cell(row.cells[0])
        add_line(row.cells[0], "No work experience records found.", bold_label=False)
    else:
        for item in rows:
            row = clone_row(table, template_row_element)
            fill_work_row(row.cells[0], item, template_cell)

    for index, paragraph in enumerate(document.paragraphs):
        if "Signature over Printed Name" in paragraph.text:
            previous = document.paragraphs[index - 1] if index > 0 else paragraph.insert_paragraph_before()
            set_paragraph_text(previous, full_name(employee) or "Employee/Applicant")
            previous.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = previous.runs[0]
            run.bold = True
            apply_run_font(run)

    document.save(output_path)
    return len(rows)


def main():
    if len(sys.argv) != 4:
        print("Usage: wes_docx.py <input-json> <output-docx> <template-docx>", file=sys.stderr)
        return 2
    input_path = Path(sys.argv[1])
    output_path = Path(sys.argv[2])
    template_path = Path(sys.argv[3])
    payload = json.loads(input_path.read_text(encoding="utf-8-sig"))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    row_count = fill_docx(payload, output_path, template_path)
    print(json.dumps({"rowCount": row_count}))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
